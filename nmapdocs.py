#!/usr/bin/env python3
"""
Nmap to Word Report Generator

- Parses Nmap scan results (preferably XML via `-oX`), with a best-effort fallback for plain text.
- Generates a penetration testing style Word document using python-docx.
- Optionally logs an entry to pentest.txt to maintain chronological findings.

Usage examples:
  python nmapdocs.py --input scan.xml --output report.docx --assessment "Network Assessment" \
      --assessor "Alice Smith" --scan-command "nmap -sV -T4 -oX scan.xml 10.0.0.0/24"

  python nmapdocs.py --input scan.txt --output report.docx --assessment "Infra PT"

Recommend creating XML with Nmap:
  nmap -sV -T4 -oX scan.xml 10.0.0.0/24
"""
from __future__ import annotations

import argparse
import datetime as dt
import ipaddress
import os
import re
import sys
import xml.etree.ElementTree as ET
from typing import Dict, List, Optional, Tuple, Any
import io

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.enum.section import WD_ORIENTATION, WD_SECTION
except ImportError as e:
    print("Missing dependency 'python-docx'. Install with: pip install -r requirements.txt", file=sys.stderr)
    raise


# -----------------------------
# Data Structures
# -----------------------------
class PortInfo:
    def __init__(self, port: int, proto: str, state: str, service: str = "", product: str = "", version: str = ""):
        self.port = port
        self.proto = proto
        self.state = state
        self.service = service
        self.product = product
        self.version = version

    def to_row(self) -> List[str]:
        svc = self.service or "-"
        prodver = " ".join([p for p in [self.product, self.version] if p]).strip() or "-"
        return [str(self.port), self.proto, self.state, svc, prodver]


class HostInfo:
    def __init__(self, ip: str, hostname: str = "", status: str = "unknown"):
        self.ip = ip
        self.hostname = hostname
        self.status = status  # up/down
        self.ports: List[PortInfo] = []

    @property
    def is_up(self) -> bool:
        return self.status.lower() == "up"

    def open_ports(self) -> List[PortInfo]:
        return [p for p in self.ports if p.state.lower() == "open"]


class ScanResult:
    def __init__(self):
        self.hosts: List[HostInfo] = []
        self.raw_text: str = ""
        self.start_time: Optional[dt.datetime] = None
        self.end_time: Optional[dt.datetime] = None
        self.nmap_version: Optional[str] = None
        self.args: Optional[str] = None

    def up_hosts(self) -> List[HostInfo]:
        return [h for h in self.hosts if h.is_up]


# -----------------------------
# Parsing Functions
# -----------------------------

def parse_nmap_xml(text: str) -> Optional[ScanResult]:
    """Parse Nmap XML output and return ScanResult, or None on failure."""
    try:
        root = ET.fromstring(text)
    except ET.ParseError:
        return None

    if root.tag != "nmaprun":
        return None

    res = ScanResult()
    res.raw_text = text
    res.nmap_version = root.attrib.get("version")
    res.args = root.attrib.get("args")

    # Times
    try:
        start = root.attrib.get("start")
        if start:
            res.start_time = dt.datetime.fromtimestamp(int(start))
    except Exception:
        pass

    # endtime in runstats/finished
    finished = root.find("runstats/finished")
    if finished is not None:
        when = finished.attrib.get("time")
        try:
            if when:
                res.end_time = dt.datetime.fromtimestamp(int(when))
        except Exception:
            pass

    for host_node in root.findall("host"):
        status = host_node.find("status").attrib.get("state", "unknown") if host_node.find("status") is not None else "unknown"

        ip = ""
        for addr in host_node.findall("address"):
            if addr.attrib.get("addrtype") == "ipv4":
                ip = addr.attrib.get("addr", "")
                break
            if not ip and addr.attrib.get("addrtype") == "ipv6":
                ip = addr.attrib.get("addr", "")
        if not ip:
            # skip hosts without IP
            continue

        hostname = ""
        hostnames = host_node.find("hostnames")
        if hostnames is not None:
            hn = hostnames.find("hostname")
            if hn is not None:
                hostname = hn.attrib.get("name", "")

        h = HostInfo(ip=ip, hostname=hostname, status=status)

        ports = host_node.find("ports")
        if ports is not None:
            for p in ports.findall("port"):
                proto = p.attrib.get("protocol", "tcp")
                try:
                    portid = int(p.attrib.get("portid", "0"))
                except ValueError:
                    portid = 0
                state_node = p.find("state")
                state = state_node.attrib.get("state", "unknown") if state_node is not None else "unknown"
                svc_node = p.find("service")
                service = svc_node.attrib.get("name", "") if svc_node is not None else ""
                product = svc_node.attrib.get("product", "") if svc_node is not None else ""
                version = svc_node.attrib.get("version", "") if svc_node is not None else ""

                h.ports.append(PortInfo(port=portid, proto=proto, state=state, service=service, product=product, version=version))

        res.hosts.append(h)

    return res


# -----------------------------
# Merge Utilities
# -----------------------------
def merge_scan_results(results: List[ScanResult]) -> ScanResult:
    """Merge multiple ScanResult objects into a single consolidated result.
    - Hosts are merged by IP address.
    - Host status: 'up' wins if any source has host up.
    - Hostname: prefer a non-empty hostname; keep first non-empty seen.
    - Ports: de-duplicate by (port, proto). Prefer richer service/product/version info when available.
    - Times: start_time = min, end_time = max when present.
    - Args/raw_text: concatenate for traceability.
    """
    merged = ScanResult()
    host_map: Dict[str, HostInfo] = {}

    # Merge metadata
    starts = []
    ends = []
    arg_snippets: List[str] = []
    raw_snippets: List[str] = []

    for res in results:
        if res.start_time:
            starts.append(res.start_time)
        if res.end_time:
            ends.append(res.end_time)
        if res.args:
            arg_snippets.append(res.args)
        if res.raw_text:
            raw_snippets.append(res.raw_text)

        for h in res.hosts:
            ip = h.ip
            if not ip:
                continue
            if ip not in host_map:
                # clone host
                new_h = HostInfo(ip=ip, hostname=h.hostname, status=h.status)
                new_h.ports = []
                host_map[ip] = new_h
            dst = host_map[ip]
            # status merge
            if h.is_up:
                dst.status = "up"
            # hostname merge
            if not dst.hostname and h.hostname:
                dst.hostname = h.hostname
            # merge ports
            existing: Dict[Tuple[int, str], PortInfo] = {(p.port, p.proto): p for p in dst.ports}
            for p in h.ports:
                key = (p.port, p.proto)
                if key not in existing:
                    # copy
                    dst.ports.append(PortInfo(p.port, p.proto, p.state, p.service, p.product, p.version))
                    existing[key] = dst.ports[-1]
                else:
                    cur = existing[key]
                    # prefer 'open' state if either is open
                    if p.state.lower() == "open":
                        cur.state = "open"
                    # fill in richer service/product/version if missing
                    if (not cur.service) and p.service:
                        cur.service = p.service
                    if (not cur.product) and p.product:
                        cur.product = p.product
                    if (not cur.version) and p.version:
                        cur.version = p.version

    # finalize
    merged.hosts = list(host_map.values())
    if starts:
        merged.start_time = min(starts)
    if ends:
        merged.end_time = max(ends)
    if arg_snippets:
        merged.args = " | ".join(arg_snippets)
    if raw_snippets:
        merged.raw_text = "\n\n".join(raw_snippets)

    return merged

def parse_nmap_text(text: str) -> ScanResult:
    """Best-effort plain text parser for normal Nmap output."""
    res = ScanResult()
    res.raw_text = text

    current_host: Optional[HostInfo] = None

    # Patterns
    host_re = re.compile(r"^Nmap scan report for (.+?)(?: \(([^\)]+)\))?$")
    host_up_re = re.compile(r"^Host is up")
    addr_re = re.compile(r"^Nmap scan report for (\d+\.\d+\.\d+\.\d+)$")

    # Port lines after header "PORT\tSTATE\tSERVICE"; also from -sV with product/version stub at end.
    port_line_re = re.compile(r"^(\d+)/(tcp|udp)\s+([a-zA-Z0-9_\-]+)\s+([\w\-\?]+)(?:\s+(.+))?$")

    for line in text.splitlines():
        line = line.rstrip()
        m = host_re.match(line)
        if m:
            # Commit previous host
            if current_host is not None:
                res.hosts.append(current_host)
            hostname = m.group(1)
            ip = m.group(2) or ""
            # If the hostname is actually an IP and ip is empty, set it
            if not ip:
                try:
                    ipaddress.ip_address(hostname)
                    ip = hostname
                    hostname = ""
                except Exception:
                    pass
            current_host = HostInfo(ip=ip, hostname=hostname, status="down")
            continue

        if current_host is not None:
            if host_up_re.match(line):
                current_host.status = "up"
                continue
            if line.startswith("PORT"):
                # header line, skip
                continue
            pm = port_line_re.match(line)
            if pm:
                port = int(pm.group(1))
                proto = pm.group(2)
                state = pm.group(3)
                service = pm.group(4)
                prodver = pm.group(5) or ""
                product = ""
                version = ""
                if prodver:
                    # naive split like "Apache httpd 2.4.57"
                    parts = prodver.split()
                    if len(parts) >= 2:
                        product = " ".join(parts[:-1])
                        version = parts[-1]
                    else:
                        product = prodver
                current_host.ports.append(PortInfo(port, proto, state, service, product, version))
                continue

    if current_host is not None:
        res.hosts.append(current_host)

    return res


# -----------------------------
# Document Generation
# -----------------------------

def set_document_defaults(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    # Make all paragraphs justified by default
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Apply common margins to the initial section
    def _set_margins(sec, inches: float = 1.0):
        sec.left_margin = Inches(inches)
        sec.right_margin = Inches(inches)
        sec.top_margin = Inches(inches)
        sec.bottom_margin = Inches(inches)

    if hasattr(doc, 'sections') and doc.sections:
        _set_margins(doc.sections[0], 1.0)

    # expose helper for other functions
    doc._set_common_margins = _set_margins  # type: ignore[attr-defined]


def add_title_page(doc: Document, assessment: str, assessor: Optional[str], scan_result: ScanResult):
    title = doc.add_paragraph()
    run = title.add_run(assessment or "Security Assessment Report")
    run.font.size = Pt(28)
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    date_str = dt.datetime.now().strftime("%Y-%m-%d")
    subtitle_run = subtitle.add_run(f"Date: {date_str}")
    subtitle_run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if assessor:
        assessor_p = doc.add_paragraph()
        assessor_run = assessor_p.add_run(f"Assessor: {assessor}")
        assessor_run.font.size = Pt(12)
        assessor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def add_executive_summary(doc: Document, scan_result: ScanResult):
    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph(
        "This assessment focuses on identifying active hosts and exposed network services discovered via Nmap. "
        "The findings below summarize reachable assets and open ports which may increase the attack surface. "
        "Further validation and risk analysis are recommended for production environments."
    )
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)


def add_scope_and_methodology(doc: Document, scan_result: ScanResult, scan_command: Optional[str]):
    doc.add_heading("Scope & Methodology", level=1)
    # Brief description
    doc.add_paragraph(
        "The following table lists the scan targets as reflected in the Nmap results, "
        "including the resolved hostname where available."
    )

    # Build table of all hosts present in results (represents provided input/targets)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 1'
    headers = ["Target (IP)", "Hostname"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    try:
        table.allow_autofit = False  # type: ignore[attr-defined]
    except Exception:
        pass
    widths = [Inches(1.8), Inches(3.8)]
    for idx, w in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = w
    # Left-align header
    for cell in table.rows[0].cells:
        if cell.paragraphs:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in cell.paragraphs[0].runs:
                r.bold = True

    if not scan_result.hosts:
        row = table.add_row().cells
        row[0].text = "-"
        row[1].text = "No targets found in results"
    else:
        for h in scan_result.hosts:
            row = table.add_row().cells
            row[0].text = h.ip or "-"
            row[1].text = h.hostname or "(no hostname)"

    doc.add_paragraph("")


def add_hosts_overview(doc: Document, scan_result: ScanResult):
    doc.add_heading("Discovered Hosts Overview", level=1)
    # Brief description under the heading
    doc.add_paragraph(
        "This section summarizes the hosts that responded during the scan window. "
        "Only systems observed as actively reachable (host state: up) are listed below, "
        "along with their resolved hostnames and IP addresses."
    )

    # Only list active hosts, numbered
    doc.add_paragraph("Active Hosts:")
    up_hosts = scan_result.up_hosts()
    if not up_hosts:
        doc.add_paragraph("No active hosts detected.")
    else:
        for h in up_hosts:
            name = h.hostname or "(no hostname)"
            p = doc.add_paragraph(f"{name} [{h.ip}]")
            # Numbered list formatting
            p.style = 'List Number'
            # Tight spacing between items
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)


def add_ports_tables(doc: Document, scan_result: ScanResult):
    doc.add_heading("Open Ports & Services", level=1)
    # Brief description under the heading
    doc.add_paragraph(
        "For each active host, this table enumerates the open network ports and associated services as identified by Nmap. "
        "Where available, product names and versions are included to assist with vulnerability assessment and patch validation."
    )

    up_hosts = scan_result.up_hosts()
    if not up_hosts:
        doc.add_paragraph("No active hosts to display.")
        return

    for host in up_hosts:
        htitle = doc.add_paragraph()
        run = htitle.add_run(f"Host: {host.hostname or host.ip} ({host.ip})")
        run.bold = True

        open_ports = host.open_ports()
        if not open_ports:
            doc.add_paragraph("No open ports detected.")
            continue

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light List Accent 1'
        hdr_cells = table.rows[0].cells
        headers = ["Port", "Proto", "State", "Service", "Product/Version"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        # Ensure consistent column widths for clean header alignment
        # Widths tuned for readability: narrow numeric columns first, then wider text columns
        col_widths = [Inches(0.9), Inches(0.9), Inches(0.9), Inches(1.8), Inches(3.0)]
        try:
            # python-docx >= 0.8.11 supports allow_autofit
            table.allow_autofit = False  # type: ignore[attr-defined]
        except Exception:
            pass
        for idx, width in enumerate(col_widths):
            for cell in table.columns[idx].cells:
                cell.width = width

        # Left-align and bold the header row
        for cell in hdr_cells:
            if cell.paragraphs:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True

        # Sort ports numerically
        for p in sorted(open_ports, key=lambda x: (x.proto, x.port)):
            row_cells = table.add_row().cells
            for i, val in enumerate(p.to_row()):
                row_cells[i].text = val

        # Add a bit of space after table
        doc.add_paragraph("")


def add_overall_scan_summary(doc: Document, scan_result: ScanResult):
    """Create a new landscape page with a bar chart summarizing number of open ports per active host.
    If matplotlib is not available, fall back to a compact table.
    """
    # Start a new section in landscape
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    # Capture original (likely portrait) dimensions
    orig_w, orig_h = sec.page_width, sec.page_height
    sec.orientation = WD_ORIENTATION.LANDSCAPE
    # Force landscape dimensions explicitly
    sec.page_width, sec.page_height = (max(orig_w, orig_h), min(orig_w, orig_h))
    # Standardized margins
    try:
        doc._set_common_margins(sec, 1.0)  # type: ignore[attr-defined]
    except Exception:
        pass

    doc.add_heading("Overall Scan Summary", level=1)
    doc.add_paragraph(
        "This summary shows, for each active system identified, the total number of open network ports. "
        "Higher counts generally indicate a larger attack surface and warrant further review."
    )

    # Prepare data
    up_hosts = scan_result.up_hosts()
    labels = [(h.hostname or h.ip or "(unknown)") for h in up_hosts]
    values = [len(h.open_ports()) for h in up_hosts]

    added_visual = False
    if up_hosts and any(values):
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt

            fig, ax = plt.subplots(figsize=(10, 4.2))
            ax.bar(range(len(values)), values, color="#4F81BD")
            ax.set_xticks(range(len(labels)))
            # Shorten long labels
            short_labels = [lbl if len(lbl) <= 18 else lbl[:17] + "…" for lbl in labels]
            ax.set_xticklabels(short_labels, rotation=30, ha="right")
            ax.set_ylabel("Open Ports")
            ax.set_xlabel("Host")
            ax.set_title("Open Ports per Active Host")
            ax.grid(axis="y", linestyle=":", alpha=0.6)

            buf = io.BytesIO()
            fig.tight_layout()
            fig.savefig(buf, format="png", dpi=200)
            plt.close(fig)
            buf.seek(0)

            # Fit image to available width
            page_width = sec.page_width
            left = sec.left_margin
            right = sec.right_margin
            available_width = page_width - left - right
            pic = doc.add_picture(buf)
            pic.width = available_width
            try:
                from docx.shared import Inches as _In
                pic.height = _In(3.2)
            except Exception:
                pass
            # Interpretation (below graph, same page)
            doc.add_paragraph(
                "Interpretation: Each bar shows the total count of open TCP/UDP ports detected on that host. "
                "Taller bars generally indicate a larger attack surface. Use together with the sensitive ports chart to prioritize hosts."
            )
            added_visual = True
        except Exception:
            added_visual = False

    if not added_visual:
        # Table fallback if matplotlib is missing or error occurred
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light List Accent 1'
        hdrs = ["Hostname/IP", "Status", "Open Ports Count"]
        for i, t in enumerate(hdrs):
            table.rows[0].cells[i].text = t
        try:
            table.allow_autofit = False  # type: ignore[attr-defined]
        except Exception:
            pass
        widths = [Inches(3.5), Inches(1.2), Inches(1.5)]
        for idx, w in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = w
        for h, v in zip(up_hosts, values):
            row = table.add_row().cells
            row[0].text = (h.hostname or h.ip or "(unknown)") + (f" [{h.ip}]" if h.hostname else "")
            row[1].text = "UP" if h.is_up else "DOWN"
            row[2].text = str(v)
        # Interpretation under table
        doc.add_paragraph(
            "Interpretation: Open ports per host indicate potential exposure. Higher numbers suggest systems that warrant focused review."
        )

    # Do not reset to portrait here; the next summary section will handle portrait reset


def add_service_summary(doc: Document, scan_result: ScanResult):
    """Create a new landscape page with a bar chart showing count of open ports by service across all active hosts.
    Adds data labels on bars. Falls back to a table if matplotlib is not available.
    """
    # New landscape section
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    # Capture original (likely portrait) dimensions
    orig_w, orig_h = sec.page_width, sec.page_height
    sec.orientation = WD_ORIENTATION.LANDSCAPE
    sec.page_width, sec.page_height = (max(orig_w, orig_h), min(orig_w, orig_h))
    # Standardized margins
    try:
        doc._set_common_margins(sec, 1.0)  # type: ignore[attr-defined]
    except Exception:
        pass

    doc.add_heading("Overall Service Exposure", level=1)
    doc.add_paragraph(
        "This chart aggregates open services discovered across all active hosts and shows how many instances of each "
        "service are exposed. Higher counts can indicate shared configurations or systemic gaps."
    )

    # Aggregate service counts
    svc_counts: Dict[str, int] = {}
    for h in scan_result.up_hosts():
        for p in h.open_ports():
            svc = (p.service or "unknown").lower()
            svc_counts[svc] = svc_counts.get(svc, 0) + 1

    added_visual = False
    if svc_counts:
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt

            # Sort by count descending and keep top 20 for readability
            items = sorted(svc_counts.items(), key=lambda kv: kv[1], reverse=True)
            top = items[:20]
            labels = [k for k, _ in top]
            values = [v for _, v in top]

            fig, ax = plt.subplots(figsize=(10, 4.2))
            y = range(len(values))
            ax.barh(y, values, color="#9BBB59")
            ax.set_yticks(y)
            # shorten very long labels
            short_labels = [lbl if len(lbl) <= 20 else lbl[:19] + "…" for lbl in labels]
            ax.set_yticklabels(short_labels)
            ax.invert_yaxis()  # highest at top
            ax.set_xlabel("Instances across hosts")
            ax.set_title("Top Exposed Services (by count)")
            ax.grid(axis="x", linestyle=":", alpha=0.6)

            # Data labels
            for i, v in enumerate(values):
                ax.text(v + max(values) * 0.01, i, str(v), va='center')

            buf = io.BytesIO()
            fig.tight_layout()
            fig.savefig(buf, format="png", dpi=200, bbox_inches='tight')
            plt.close(fig)
            buf.seek(0)

            page_width = sec.page_width
            left = sec.left_margin
            right = sec.right_margin
            available_width = page_width - left - right
            doc.add_picture(buf, width=available_width)
            # Interpretation (below graph, same page)
            doc.add_paragraph(
                "Interpretation: Each bar shows how many instances of a specific service were observed across all active hosts. "
                "Services with higher counts indicate broader exposure and should be prioritized for hardening/restriction."
            )
            added_visual = True
        except Exception:
            added_visual = False

    if not added_visual:
        # Table fallback
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light List Accent 1'
        hdrs = ["Service", "Instances"]
        for i, t in enumerate(hdrs):
            table.rows[0].cells[i].text = t
        try:
            table.allow_autofit = False  # type: ignore[attr-defined]
        except Exception:
            pass
        widths = [Inches(4.0), Inches(1.5)]
        for idx, w in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = w
        for svc, cnt in sorted(svc_counts.items(), key=lambda kv: kv[1], reverse=True)[:25]:
            row = table.add_row().cells
            row[0].text = svc
            row[1].text = str(cnt)
        doc.add_paragraph(
            "Interpretation: Table lists the most common services exposed across the environment; higher counts imply wider exposure."
        )

    # Do not reset to portrait here; a following section handles the portrait reset


def add_sensitive_ports_by_host(doc: Document, scan_result: ScanResult):
    """Landscape section: stacked bar chart of counts of sensitive services per host."""
    # Define sensitive mapping similar to findings
    sensitive = {
        445: 'smb', 3389: 'rdp', 1433: 'mssql', 5985: 'winrm', 5986: 'winrm', 22: 'ssh',
        6379: 'redis', 5900: 'vnc', 1521: 'oracle', 3306: 'mysql', 5432: 'postgres', 80: 'http', 23: 'telnet'
    }

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    orig_w, orig_h = sec.page_width, sec.page_height
    sec.orientation = WD_ORIENTATION.LANDSCAPE
    sec.page_width, sec.page_height = (max(orig_w, orig_h), min(orig_w, orig_h))
    try:
        doc._set_common_margins(sec, 1.0)  # type: ignore[attr-defined]
    except Exception:
        pass

    doc.add_heading("Sensitive Ports per Host", level=1)
    doc.add_paragraph(
        "Stacked bar chart showing, for each active host, the number of exposed sensitive services (e.g., SMB, RDP)."
    )

    up_hosts = scan_result.up_hosts()
    labels = [(h.hostname or h.ip or "(unknown)") for h in up_hosts]
    # Build category counts per host
    categories = sorted(set(sensitive.values()))
    data = {cat: [0]*len(up_hosts) for cat in categories}
    for idx, h in enumerate(up_hosts):
        for p in h.open_ports():
            if p.port in sensitive:
                data[sensitive[p.port]][idx] += 1

    # Only draw if there is any data
    has_data = any(sum(vals) > 0 for vals in data.values())
    added = False
    if has_data:
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt

            fig, ax = plt.subplots(figsize=(10, 5))
            bottoms = [0]*len(up_hosts)
            colors = ["#4F81BD", "#C0504D", "#9BBB59", "#8064A2", "#4BACC6", "#F79646", "#7F7F7F", "#948A54", "#1F497D", "#E46C0A", "#31859B", "#00B0F0", "#92D050"]
            for i, cat in enumerate(categories):
                vals = data[cat]
                ax.bar(range(len(labels)), vals, bottom=bottoms, label=cat, color=colors[i % len(colors)])
                bottoms = [b+v for b, v in zip(bottoms, vals)]
            ax.set_xticks(range(len(labels)))
            short_labels = [lbl if len(lbl) <= 18 else lbl[:17] + "…" for lbl in labels]
            ax.set_xticklabels(short_labels, rotation=30, ha='right')
            ax.set_ylabel('Sensitive Ports Count')
            ax.set_xlabel('Host')
            ax.set_title('Sensitive Ports per Host (stacked by service)')
            ax.legend(loc='upper right', ncol=3, fontsize=8)
            ax.grid(axis='y', linestyle=':', alpha=0.6)

            buf = io.BytesIO()
            fig.tight_layout()
            fig.savefig(buf, format='png', dpi=200)
            plt.close(fig)
            buf.seek(0)
            page_width = sec.page_width
            left = sec.left_margin
            right = sec.right_margin
            available_width = page_width - left - right
            doc.add_picture(buf, width=available_width)
            doc.add_paragraph(
                "Interpretation: Stacked colors represent different sensitive services (e.g., SMB, RDP). "
                "The total height equals the number of sensitive ports open on that host, helping identify priority systems."
            )
            added = True
        except Exception:
            added = False

    if not added:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light List Accent 1'
        hdr = ["Hostname/IP", "Sensitive Services", "Total Sensitive Ports"]
        for i, htxt in enumerate(hdr):
            table.rows[0].cells[i].text = htxt
        try:
            table.allow_autofit = False  # type: ignore[attr-defined]
        except Exception:
            pass
        widths = [Inches(3.5), Inches(2.5), Inches(1.2)]
        for idx, w in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = w
        for i, h in enumerate(up_hosts):
            total = sum(data[cat][i] for cat in categories)
            present = [cat for cat in categories if data[cat][i] > 0]
            row = table.add_row().cells
            row[0].text = (h.hostname or h.ip or "(unknown)") + (f" [{h.ip}]" if h.hostname else "")
            row[1].text = ", ".join(present) if present else "-"
            row[2].text = str(total)
        doc.add_paragraph(
            "Interpretation: Table summarizes which sensitive services each host exposes and the total count per host."
        )

    # don't reset orientation here; a subsequent section will handle portrait


def add_top_ports_summary(doc: Document, scan_result: ScanResult):
    """Landscape section: horizontal bar chart of number of hosts exposing each port (top N)."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    orig_w, orig_h = sec.page_width, sec.page_height
    sec.orientation = WD_ORIENTATION.LANDSCAPE
    sec.page_width, sec.page_height = (max(orig_w, orig_h), min(orig_w, orig_h))
    try:
        doc._set_common_margins(sec, 1.0)  # type: ignore[attr-defined]
    except Exception:
        pass

    doc.add_heading("Top Ports Across Network", level=1)
    doc.add_paragraph(
        "Horizontal bar chart showing the number of hosts exposing each open port (top 20)."
    )

    counts: Dict[int, int] = {}
    for h in scan_result.up_hosts():
        # de-duplicate ports per host to count 'number of hosts exposing'
        seen = set()
        for p in h.open_ports():
            key = p.port
            if key in seen:
                continue
            seen.add(key)
            counts[key] = counts.get(key, 0) + 1

    added = False
    if counts:
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt

            items = sorted(counts.items(), key=lambda kv: kv[1], reverse=True)[:20]
            ports = [str(k) for k, _ in items]
            values = [v for _, v in items]

            fig, ax = plt.subplots(figsize=(10, 5))
            y = range(len(values))
            ax.barh(y, values, color='#4BACC6')
            ax.set_yticks(y)
            ax.set_yticklabels(ports)
            ax.invert_yaxis()
            ax.set_xlabel('Number of Hosts')
            ax.set_title('Top Ports Exposed Across Hosts')
            ax.grid(axis='x', linestyle=':', alpha=0.6)
            for i, v in enumerate(values):
                ax.text(v + max(values) * 0.01, i, str(v), va='center')
            buf = io.BytesIO()
            fig.tight_layout()
            fig.savefig(buf, format='png', dpi=200, bbox_inches='tight')
            plt.close(fig)
            buf.seek(0)
            page_width = sec.page_width
            left = sec.left_margin
            right = sec.right_margin
            available_width = page_width - left - right
            doc.add_picture(buf, width=available_width)
            # Interpretation (below graph, same page)
            doc.add_paragraph(
                "Interpretation: Each bar shows how many hosts expose the given TCP/UDP port. "
                "Focus remediation and controls on ports with the highest host counts first."
            )
            added = True
        except Exception:
            added = False

    if not added:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light List Accent 1'
        hdrs = ["Port", "Hosts Exposing"]
        for i, t in enumerate(hdrs):
            table.rows[0].cells[i].text = t
        for k, v in sorted(counts.items(), key=lambda kv: kv[1], reverse=True)[:25]:
            row = table.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)
        doc.add_paragraph(
            "Interpretation: Table ranks ports by the number of hosts exposing them; higher counts indicate broader exposure."
        )

    # Resume portrait for remainder of the document
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec2.orientation = WD_ORIENTATION.PORTRAIT
    sec2.page_width, sec2.page_height = (min(orig_w, orig_h), max(orig_w, orig_h))
    try:
        doc._set_common_margins(sec2, 1.0)  # type: ignore[attr-defined]
    except Exception:
        pass

def add_findings_notes(doc: Document, scan_result: ScanResult):
    doc.add_heading("Findings & Recommendations", level=1)
    doc.add_paragraph(
        "This section highlights notable exposures derived from the scan results to help prioritize remediation efforts."
    )

    # Sensitive ports commonly targeted by attackers (non-exhaustive)
    SENSITIVE_PORTS = {
        20, 21, 22, 23, 25, 53, 69, 80, 110, 111, 123, 137, 139, 143, 161, 389,
        445, 512, 513, 514, 873, 902, 1080, 1433, 1521, 2049, 2375, 2376, 2379, 2380,
        3000, 3306, 3389, 4333, 5000, 5432, 5601, 5900, 5985, 5986, 6379, 6667, 7001,
        8000, 8080, 8081, 8088, 8443, 8888, 9000, 9090, 9200, 9418, 11211, 15672, 27017
    }

    # Collect sensitive port exposures
    sensitive_items = []
    for h in scan_result.up_hosts():
        for p in h.open_ports():
            if p.port in SENSITIVE_PORTS:
                svc = p.service or "unknown"
                sensitive_items.append((p.port, p.proto, svc, h.hostname or "(no hostname)", h.ip))

    doc.add_heading("Sensitive Ports", level=2)
    doc.add_paragraph(
        "The following entries show services exposed on ports commonly associated with elevated risk (remote administration, "
        "legacy or unauthenticated protocols, databases, or management interfaces)."
    )
    if not sensitive_items:
        doc.add_paragraph("None identified.")
    else:
        # Sort by port then host
        data = sorted(sensitive_items, key=lambda x: (x[0], x[3], x[4]))
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light List Accent 1'
        hdr = ["Port", "Proto", "Service", "Hostname", "IP"]
        for i, htxt in enumerate(hdr):
            table.rows[0].cells[i].text = htxt
        # Set column widths
        try:
            table.allow_autofit = False  # type: ignore[attr-defined]
        except Exception:
            pass
        widths = [Inches(0.9), Inches(0.9), Inches(1.8), Inches(2.2), Inches(1.6)]
        for idx, w in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = w
        # Left-align header
        for cell in table.rows[0].cells:
            if cell.paragraphs:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in cell.paragraphs[0].runs:
                    r.bold = True
        # Rows
        for port, proto, svc, host, ip in data:
            row = table.add_row().cells
            row[0].text = str(port)
            row[1].text = proto
            row[2].text = svc or "-"
            row[3].text = host
            row[4].text = ip
        doc.add_paragraph("")

    # Collect services where a clear version was detected by Nmap (-sV)
    versioned_items = []
    for h in scan_result.up_hosts():
        for p in h.open_ports():
            # Require an explicit version string; product name alone is insufficient
            if p.version and p.version.strip():
                prodver = " ".join([s for s in [p.product, p.version] if s]).strip()
                svc = p.service or "unknown"
                versioned_items.append((h.hostname or "(no hostname)", h.ip, p.port, p.proto, svc, prodver))

    doc.add_heading("Exposed Services with Version", level=2)
    doc.add_paragraph(
        "These services disclose a product and/or version. Precise version data enables targeted exploitation and should be "
        "minimized where possible (e.g., limit banners, enforce TLS, or restrict exposure)."
    )
    if not versioned_items:
        doc.add_paragraph("None identified.")
    else:
        # Sort by service name then version, then host
        data = sorted(versioned_items, key=lambda x: (x[4] or "", x[5] or "", x[0], x[1]))
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light List Accent 1'
        hdr = ["IP", "Port", "Proto", "Service", "Product/Version"]
        for i, htxt in enumerate(hdr):
            table.rows[0].cells[i].text = htxt
        # Set column widths
        try:
            table.allow_autofit = False  # type: ignore[attr-defined]
        except Exception:
            pass
        # Allocate more width to Product/Version to avoid wrapping; rebalance others
        widths = [Inches(1.5), Inches(0.8), Inches(0.8), Inches(1.8), Inches(2.8)]
        for idx, w in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = w
        # Left-align header
        for cell in table.rows[0].cells:
            if cell.paragraphs:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in cell.paragraphs[0].runs:
                    r.bold = True
        # Rows
        for host, ip, port, proto, svc, prodver in data:
            row = table.add_row().cells
            # IP (left)
            row[0].text = ip
            if row[0].paragraphs:
                row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Port (center)
            row[1].text = str(port)
            if row[1].paragraphs:
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Proto (center)
            row[2].text = proto
            if row[2].paragraphs:
                row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Service (left)
            row[3].text = svc or "-"
            if row[3].paragraphs:
                row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Product/Version (left)
            row[4].text = prodver or "-"
            if row[4].paragraphs:
                row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph("")


def add_appendix_raw_output(doc: Document, raw_text: str):
    doc.add_heading("Appendix: Full Nmap Output", level=1)
    if not raw_text:
        doc.add_paragraph("No raw output available.")
        return
    # Use a basic monospace style by setting font; python-docx doesn't ship with code style by default
    p = doc.add_paragraph()
    run = p.add_run(raw_text)
    run.font.name = "Consolas"
    try:
        # ensure correct rPr setting for East Asia fonts as well
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    except Exception:
        pass


# -----------------------------
# Appendix: Risks (based on sensitive ports)
# -----------------------------
def add_appendix_risks(doc: Document, scan_result: ScanResult):
    # Start Appendix on a new page
    doc.add_page_break()
    doc.add_heading("Appendix: Risks", level=1)
    # Determine which sensitive-port categories are present and provide 4–5 management-friendly lines per risk
    SENSITIVE_DETAILS = {
        3389: [
            "Remote Desktop (RDP) was found accessible from the network.",
            "This enables direct logon attempts to Windows servers and workstations.",
            "If credentials are guessed, stolen, or reused, an attacker can take control of the system.",
            "From there they can move to other systems and access business data.",
            "Limiting exposure and enforcing strong MFA significantly reduces this risk.",
        ],
        445: [
            "Microsoft file sharing (SMB) is reachable from the network.",
            "SMB is frequently targeted by ransomware and credential attacks.",
            "Compromise can allow attackers to read or encrypt files and spread to other machines.",
            "Restricting SMB to internal segments and hardening authentication reduces impact.",
            "Legacy protocols and unused shares should be disabled to shrink the attack surface.",
        ],
        1433: [
            "Microsoft SQL Server appears exposed to the network.",
            "Databases often store sensitive business information and credentials.",
            "If weak passwords or known flaws exist, attackers may steal or alter data.",
            "Restricting access and enforcing least-privilege accounts helps protect data.",
            "Regular patching and monitoring are essential to keep this risk low.",
        ],
        22: [
            "Secure Shell (SSH) is open for remote access.",
            "While encrypted, it is a common target for password-guessing and key misuse.",
            "If accessed, attackers can run commands and pivot to other systems.",
            "Use strong keys, disable passwords, and restrict who can connect to reduce risk.",
            "Network allow‑lists or VPN access provide additional protection.",
        ],
        80: [
            "A web service is available over HTTP without encryption.",
            "Information and credentials sent over HTTP can be read by anyone on the path.",
            "Unencrypted services are often older and may contain known weaknesses.",
            "Migrating to HTTPS and keeping platforms updated mitigates this exposure.",
            "Consider limiting access to administrative interfaces.",
        ],
        5985: [
            "Windows Remote Management (WinRM) over HTTP is exposed.",
            "This allows remote administration and, if misconfigured, remote command execution.",
            "Attackers with valid credentials can automate actions across multiple hosts.",
            "Enforce HTTPS, restrict who can use WinRM, and monitor for unusual activity.",
            "Disable where not required to reduce the attack surface.",
        ],
        5986: [
            "WinRM over HTTPS is available for remote administration.",
            "Although encrypted, it still relies on strong authentication and configuration.",
            "If accounts are compromised, an attacker gains remote control capabilities.",
            "Restrict access, enforce MFA, and log admin actions to reduce risk.",
            "Disable or limit to management networks where possible.",
        ],
        3306: [
            "A MySQL database service is reachable from the network.",
            "Databases may expose critical customer and business records if accessed.",
            "Weak accounts or outdated software increase the chance of data theft.",
            "Limit access to specific application servers and enforce strong credentials.",
            "Regular updates and backups help contain impact if an incident occurs.",
        ],
        5432: [
            "A PostgreSQL database service is accessible from the network.",
            "Unauthorized access could allow reading or changing important data.",
            "Risk increases with weak passwords or missing updates.",
            "Restrict exposure, use least-privilege roles, and keep the service patched.",
            "Enable monitoring and alerts for suspicious queries or connections.",
        ],
        1521: [
            "An Oracle database listener is reachable from the network.",
            "Listeners may reveal information and are targeted for authentication bypass.",
            "Compromise can lead to data loss or operational disruption.",
            "Restrict access to known application hosts and enforce strong controls.",
            "Review Oracle security guides to harden the listener configuration.",
        ],
        6379: [
            "A Redis service is accessible.",
            "Default or unauthenticated access can allow an attacker to read and change data.",
            "In some cases, it can even be abused to run code on the server.",
            "Bind Redis to localhost or trusted networks and require authentication.",
            "Disable dangerous commands where appropriate and keep software updated.",
        ],
        5900: [
            "VNC remote desktop is exposed to the network.",
            "VNC often uses weak or no encryption and can be brute‑forced.",
            "If accessed, attackers gain screen and keyboard control of the system.",
            "Restrict access, require strong credentials, and consider VPN-only use.",
            "Disable when not needed to reduce exposure.",
        ],
        23: [
            "Telnet, an older remote access protocol, is exposed.",
            "Telnet sends information in clear text, which can be intercepted.",
            "Attackers can capture passwords and issue commands as if they were the user.",
            "Replace Telnet with SSH and block Telnet at network boundaries.",
            "Remove legacy services that are no longer necessary.",
        ],
    }

    present_ports = set()
    for h in scan_result.up_hosts():
        for p in h.open_ports():
            if p.port in SENSITIVE_DETAILS:
                present_ports.add(p.port)

    if not present_ports:
        doc.add_paragraph("No high-risk service exposures were identified among the sensitive ports assessed.")
        return

    # Order by perceived severity (SMB/RDP first), then by port number
    severity_order = [445, 3389, 1433, 5985, 5986, 22, 6379, 5900, 1521, 3306, 5432, 80, 23]
    def _sev_key(port: int) -> tuple:
        return (severity_order.index(port) if port in severity_order else len(severity_order), port)

    for port in sorted(present_ports, key=_sev_key):
        details = SENSITIVE_DETAILS.get(port)
        if not details:
            details = [
                f"Service on port {port} is accessible from the network. ",
                "If accounts or configurations are weak, attackers may gain access. ",
                "Successful access could result in data exposure or system misuse. ",
                "Limit exposure, enforce strong authentication, and monitor closely.",
            ]
        # Port label (bold)
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(f"Port {port} Risk Summary")
        title_run.bold = True
        # Single explanatory paragraph (4–5 sentences) with proper spacing
        para_text = " ".join(s.strip() for s in details[:5]).strip()
        doc.add_paragraph(para_text)
        # Affected systems for this port
        affected = []
        for h in scan_result.up_hosts():
            for p in h.open_ports():
                if p.port == port:
                    affected.append((h.hostname or "(no hostname)", h.ip))
                    break
        if affected:
            aff_title = doc.add_paragraph()
            aff_title_run = aff_title.add_run("Affected systems:")
            aff_title_run.bold = True
            for host, ip in sorted(affected, key=lambda x: (x[0], x[1])):
                doc.add_paragraph(f"- {host} [{ip}]", style='List Bullet')
        doc.add_paragraph("")


# Pentest log feature removed per user request.


# -----------------------------
# Main
# -----------------------------

def load_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def build_report(input_paths: List[str], output_path: str, assessment: str, assessor: Optional[str], scan_command: Optional[str]) -> str:
    """Build a report from one or more Nmap outputs. Multiple inputs are merged (e.g., TCP + UDP)."""
    parsed_results: List[ScanResult] = []
    for pth in input_paths:
        raw = load_text(pth)
        parsed = parse_nmap_xml(raw)
        if parsed is None:
            parsed = parse_nmap_text(raw)
        # Track the original args/file as context
        if not parsed.args:
            parsed.args = f"(input: {os.path.basename(pth)})"
        parsed_results.append(parsed)

    # Merge results if more than one
    parsed = merge_scan_results(parsed_results) if len(parsed_results) > 1 else parsed_results[0]

    # Build document
    doc = Document()
    set_document_defaults(doc)

    add_title_page(doc, assessment=assessment, assessor=assessor, scan_result=parsed)
    add_executive_summary(doc, scan_result=parsed)
    add_scope_and_methodology(doc, scan_result=parsed, scan_command=scan_command or parsed.args)
    add_hosts_overview(doc, scan_result=parsed)
    add_ports_tables(doc, scan_result=parsed)
    add_overall_scan_summary(doc, scan_result=parsed)
    add_service_summary(doc, scan_result=parsed)
    add_sensitive_ports_by_host(doc, scan_result=parsed)
    add_top_ports_summary(doc, scan_result=parsed)
    add_findings_notes(doc, scan_result=parsed)
    # Removed Appendix section as requested
    add_appendix_risks(doc, scan_result=parsed)

    # Ensure dir exists
    os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
    doc.save(output_path)

    return output_path


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Convert Nmap results to a penetration testing style Word report.")
    parser.add_argument("--input", nargs='+', required=True, help="One or more Nmap result files to include (XML preferred; text supported)")
    parser.add_argument("--output", required=True, help="Path to the output .docx report")
    parser.add_argument("--assessment", required=False, default="Security Assessment", help="Assessment title for the report")
    parser.add_argument("--assessor", required=False, default=None, help="Name of the assessor/author")
    parser.add_argument("--scan-command", required=False, default=None, help="Nmap command used (for documentation)")

    args = parser.parse_args(argv)

    # Collect input paths (single flag may include multiple files)
    input_paths: List[str] = list(args.input)

    # Validate paths
    missing = [p for p in input_paths if not os.path.exists(p)]
    if missing:
        print(f"Input file(s) not found: {', '.join(missing)}", file=sys.stderr)
        return 2

    try:
        out = build_report(
            input_paths=input_paths,
            output_path=args.output,
            assessment=args.assessment,
            assessor=args.assessor,
            scan_command=args.scan_command,
        )
        print(f"Report generated: {out}")
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
